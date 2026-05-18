from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "NexusFlow_Final_Submission.docx"
ARCH = ROOT / "docs" / "nexus-architecture.png"

LIVE_URL = "https://nexus-flow-atomberg.vercel.app/"
REPO_URL = "https://github.com/SM33-07/NexusFlow"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="CBD5E1", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color="111827", size=10.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(r_pr)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def paragraph(doc, text="", size=10.5, color="334155", bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def heading(doc, text, level=1):
    p = paragraph(doc, text, size=16 if level == 1 else 13, color="064E3B" if level == 1 else "1E40AF", bold=True, after=8, before=14)
    return p


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    labels = [("3", "Role journeys"), ("100%", "Weightage guard"), ("Live", "Hosted demo"), ("CSV", "Admin export")]
    for index, (value, label) in enumerate(labels):
        cell = table.cell(0, index)
        set_cell_shading(cell, "ECFDF5" if index % 2 == 0 else "EFF6FF")
        set_cell_border(cell, "A7F3D0" if index % 2 == 0 else "BFDBFE")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        r.bold = True
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor.from_string("064E3B")
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label.upper())
        r2.bold = True
        r2.font.name = "Arial"
        r2._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor.from_string("64748B")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)

    paragraph(doc, "ATOMQUEST HACKATHON SUBMISSION", size=9, color="059669", bold=True, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    title = paragraph(doc, "NexusFlow", size=34, color="0F172A", bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_before = Pt(4)
    paragraph(
        doc,
        "Interactive enterprise goal canvas with role-based dashboards, approvals, quarterly tracking, audit logs, analytics, and database-backed persistence.",
        size=13,
        color="334155",
        after=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_metric_strip(doc)

    heading(doc, "1. Live / Hosted Demo URL")
    p = paragraph(doc, "Portal: ", size=11, color="0F172A", bold=True, after=2)
    add_hyperlink(p, LIVE_URL, LIVE_URL)

    heading(doc, "2. Source Code Repository")
    p = paragraph(doc, "Repository: ", size=11, color="0F172A", bold=True, after=2)
    add_hyperlink(p, REPO_URL, REPO_URL)

    heading(doc, "3. Login Credentials")
    creds = doc.add_table(rows=1, cols=3)
    creds.autofit = False
    headers = ["Role", "Email", "Password"]
    for i, header in enumerate(headers):
        cell = creds.cell(0, i)
        set_cell_shading(cell, "064E3B")
        set_cell_border(cell, "065F46")
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=10)

    rows = [
        ("Employee", "employee@nexus.demo", "NexusEmp@2026"),
        ("Manager", "manager@nexus.demo", "NexusMgr@2026"),
        ("Admin / HR", "admin@nexus.demo", "NexusAdmin@2026"),
    ]
    for role, email, password in rows:
        cells = creds.add_row().cells
        for i, value in enumerate((role, email, password)):
            set_cell_shading(cells[i], "F8FAFC")
            set_cell_border(cells[i])
            set_cell_text(cells[i], value, bold=(i == 0), color="0F172A", size=10)

    heading(doc, "4. Suggested Demo Script")
    steps = [
        ("Employee journey", "Show the visual goal canvas, AI-assisted goal creation, 100% weightage validation, and quarterly actual logging."),
        ("Manager journey", "Review direct-report goals, approve or return goals, observe approval locks, log check-ins, and export CSV."),
        ("Admin / HR journey", "Open governance analytics, push shared KPIs, inspect audit and notification logs, and evaluate escalations."),
    ]
    for label, detail in steps:
        p = paragraph(doc, f"{label}: ", size=10.5, color="0F172A", bold=True, after=2)
        run = p.add_run(detail)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string("475569")

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    heading(doc, "5. Architecture Diagram")
    paragraph(doc, "System design: browser users, Next.js App Router/API routes, signed cookie auth, Supabase Postgres, optional OpenAI goal generation, and integration logs.", size=10.5, color="475569", after=10)
    doc.add_picture(str(ARCH), width=Inches(7.35))

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("NexusFlow | AtomQuest Hackathon Submission")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("64748B")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
